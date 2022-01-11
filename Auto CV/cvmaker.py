import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from hyperlink import add_hyperlink
from os.path import abspath, dirname, join


lines = []
print("Paste job description here and type exit when finished:")
while True:
    line = input()
    if line != 'exit':
        lines.append(line)
    else:
        break
job_desc = ('\n'.join(lines)).lower()
excluded = docx.Document()
excluded.save("Excluded.docx")
document = docx.Document()
document.save("Michael Choi.docx")

# Style for Heading 1
style = document.styles['Heading 1'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Calibri'
style.size = Pt(14)
# Font change for heading and titles bugged in docx module asciiTheme also needed to be changed to  Arial.
rFonts = style.element.rPr.rFonts
rFonts.set(qn('w:asciiTheme'), 'Calibri')
paragraph_format = document.styles['Heading 1'].paragraph_format
paragraph_format.space_before = Pt(0)

# Style for Heading 2
style = document.styles['Heading 2'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Calibri'
style.size = Pt(12)
rFonts = style.element.rPr.rFonts
rFonts.set(qn('w:asciiTheme'), 'Calibri')

# Style for Normal paragraph font
style = document.styles['Normal'].font
style.name = 'Calibri'
style.size = Pt(11)

# Style used for information section under heading
styles = document.styles
style = styles.add_style('info', WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles['Normal']
style = document.styles['info'].font
style.name = 'Calibri'
style.size = Pt(10)
paragraph_format = document.styles['info'].paragraph_format
paragraph_format.space_after = Pt(0)

# Style for excluded paragraph font
style = excluded.styles['Normal'].font
style.name = 'Calibri'
style.size = Pt(12)


class Format:
    def __init__(self, name, github):
        self.name = name
        self.github = github
        # Name and link to GitHub front page is centered
        document.add_heading(name, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        email = document.add_paragraph("Email: ", style='info')
        add_hyperlink(email, "kinlonchoi@gmail.com", "kinlonchoi@gmail.com")
        document.add_paragraph("Mobile: 07712636191", style='info').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = document.add_paragraph("GitHub: ", style='info')
        add_hyperlink(paragraph, github, github)
        email.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # The following three functions are made to improve readability of code
    # This is the font used for titles
    def title(self, header):
        document.add_heading(header, level=2)

    # This is the font used for paragraphs
    def para(self, sentence):
        document.add_paragraph(sentence)

    # This is the font used for bullet points
    def bullet(self, sentence):
        document.add_paragraph(sentence, style='List Bullet')

    # Scans job description for keywords
    def scan(self, words_lists):
        for items in words_lists:
            for k, v in items:
                if k in ('python', 'js', 'javascript', 'back', 'full'):
                    p = document.add_paragraph(v, style='List Bullet')
                    add_hyperlink(p,
                                  "CV Automation tool.",
                                  "https://github.com/KinLonChoi/Python-Projects")
                    break
                elif k in job_desc:
                    self.bullet(v)
                    break
                elif k in list(items)[-1]:
                    excluded.add_paragraph(v, style='List Bullet')
                else:
                    continue


cv = Format("Michael Choi", "https://github.com/KinLonChoi")
cv.para("")
google_cert = document.add_paragraph("I am a highly motivated individual with over five years of experience in the science industry looking for a"
                                      " career change to pursue my passion for data analytics. I am passionate about data and its application and"
                                      " have recently earned the ")

add_hyperlink(google_cert, "Google Data Analytics Professional Certificate.", "https://www.coursera.org/account/accomplishments/specialization/certificate/LLK5BTDKPCDC")

# Skills section add dictionary definition for each key words(search terms) with value(skills) to add as bullet point
cv.title("Skills")


# Words to be searched for in job description
code = dict.fromkeys(['python', 'R', 'full'],
                     "Proficient with using Python and basic skills in R. Project: ")
# Some search terms are shortened to match variations of words e.g. analy will match analytical, analysis etc.
sql = dict.fromkeys(['sql', 'data', 'dbms', 'analy'],
                    "Knowledge of the use of relational databases in SQL and its advanced functions.")

algo = dict.fromkeys(['tableau', 'visual'],
                     "Data visualisation using Tableau")

time = dict.fromkeys(['spreadsheet', 'excel', 'sheets', 'google', 'vba', 'macro'],
                     "Advanced skills in excel/google sheets and its functions used in data cleaning (VBA & Macros).")

report = dict.fromkeys(['commun', 'present', 'audience'],
                       "Communication skills: Can present to an audience of varying levels of knowledge.")

skill_list = [x.items() for x in (code, sql, algo, time, report)]

cv.scan(skill_list)

# Employment section same as before add dictionary definitions for search terms.
cv.title("Employment")
cv.para("Sept 2016 – Sept 2021               Tate & Lyle PLC	             Laboratory Analyst")

special = dict.fromkeys(['detail', 'standard', 'require', 'product'],
                        "Excellent attention to detail in ensuring products meets the required standards.")

design = dict.fromkeys(['strat' 'design', 'project'],
                       "Implemented and designed the appropriate analytic strategies for unique projects.")

collab = dict.fromkeys(['collab', 'team', 'improve', 'meet'],
                       "Collaborated with other departments and developed continuous improvement strategies.")

employment_list_1 = [x.items() for x in (special, design, collab)]

cv.scan(employment_list_1)

cv.para("Key achievements:")
cv.bullet("Automation of report process that saves 6 hours every week.")
cv.bullet("Created pivot tables that extract data from the SAP database for annual reports.")

cv.para("Feb 2016 – Sept 2016                Tate & Lyle PLC                      Research scientist")

plan = dict.fromkeys(['organ', 'improve', 'projects', 'plan', 'continuous'],
                     "Organised new product development and continuous improvement (CI) projects.")

meet = dict.fromkeys(['meet', 'progress', 'team'],
                     "Communicate results to the team and provide insights followed by recommendations.")

busy = dict.fromkeys(['prior', 'time', 'effici', 'busy'],
                     "Prioritised and allocated time efficiently whilst working on several projects simultaneously.")

employment_list_2 = [x.items() for x in (plan, meet, busy)]

cv.scan(employment_list_2)
cv.para("Key achievements:")
cv.bullet("Lead the CI program that leads to the safe removal of over 100 chemicals.")

# Education and Certification section
cv.title("Education and Certification")
date = document.add_paragraph("06/01/2022		", style='Normal')
add_hyperlink(date, "Google Data Analytics Professional Certificate by Google (Coursera)",
              "https://www.coursera.org/account/accomplishments/certificate/ZCNSEFJ3GN6V")

date = document.add_paragraph("07/11/2021		", style='Normal')
add_hyperlink(date, "SQL for Data Science by University of California, Davis (Coursera)",
              "https://www.coursera.org/account/accomplishments/certificate/SYPGFMFPK3YA")

date = document.add_paragraph("04/11/2021		", style='Normal')
add_hyperlink(date, "Python for Everybody by University of Michigan (Coursera)",
              "https://www.coursera.org/account/accomplishments/specialization/certificate/DK3WDTXK4ND4")

cv.para("")

cv.para("2011-2015		MChem (Hons) in Chemistry (2:1) 		                        					University of Leicester")

cv.para("2003-2011		The Bromfords School, Essex 								"
        "               A-Levels in Chemistry (B), Biology (B), and Physics (B) 				              "
        "               10 GCSEs at A*-C including Chemistry, Maths, and English")

# This will save file in same file directory as python file
document.save(join(dirname(abspath(__file__)), "Michael Choi.docx"))
excluded.save(join(dirname(abspath(__file__)), "Excluded.docx"))
