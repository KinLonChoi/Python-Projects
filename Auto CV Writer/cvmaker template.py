from docx.shared import Pt
import docx
from hyperlink import add_hyperlink
from docx.oxml.ns import qn
from os.path import abspath, dirname, join

lines = []
print("Paste job description here and type exit when finished:")
while True:
    line = input()
    if line != 'exit':
        lines.append(line)
    else:
        break
job_desc = ('\n'.join(lines)).lower()

document = docx.Document()

####################################CHANGE NAME####################################
document.save("NAME.docx")

# Style for Heading 1
style = document.styles['Heading 1'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Arial'
style.size = Pt(14)
# Font change for heading and titles bugged in docx module asciiTheme also needed to be changed to  Arial.
rFonts = style.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")

# Style for Heading 2
style = document.styles['Heading 2'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Arial'
style.size = Pt(13)
rFonts = style.element.rPr.rFonts
rFonts.set(qn("w:asciiTheme"), "Arial")


# Style for Normal paragraph font
style = document.styles['Normal'].font
style.name = 'Arial'
style.size = Pt(12)


class Tools:
    def __init__(self, name, github):
        self.name = name
        self.github = github
        document.add_heading(name, level=1).alignment = 1
        p = document.add_paragraph("GitHub: ")
        add_hyperlink(p, github, github)
        p.alignment = 1

# The following three functions are made to improve readability of code
# This is the font used for titles
    def title(self, header):
        document.add_heading(header, level=2)

# This is the font used for paragraphs
    def para(self, sentence):
        document.add_paragraph(sentence)

# This is the font used for bullet points
    def bullet(self, sentence):
        document.add_paragraph(sentence, style='List Bullet')

# Scans job description for keywords
    def scan(self, words_lists):
        for items in words_lists:
            for k, v in items:
                if k in job_desc:
                    self.bullet(v)
                    break
                else:
                    continue

####################################ADD NAME AND GITHUB WEBSITE####################################
cv = Tools("NAME", "GITHUB WEBSITE")
####################################ADD INTRODUCTION####################################
cv.para("INTRODUCTION")


# Skills section add dictionary definition for each key words(search terms) with value(skills) to add as bullet point
cv.title("Skills")

site = document.add_paragraph("GitHub website:", style='List Bullet')
####################################ADD NAME AND GITHUB WEBSITE####################################
add_hyperlink(site, "GITHUB WEBSITE", "GITHUB WEBSITE")
################### Words to be searched for in job description####################################
skill1 = dict.fromkeys(['word1', 'word2'],
                       ##################### if term is found print above as bullet point##########
                       "sentence of skill and example")

skill2 = dict.fromkeys(['word1', 'word2'],
                       "sentence of skill and example")


skill_list = [x.items() for x in (skill1, skill2)]

cv.scan(skill_list)


# Employment section same as before add dictionary definitions for search terms.
cv.title("Employment")
cv.para("word dates               company	             job title")

employment_1_1 = dict.fromkeys(['word1', 'word2'],
                            "sentence of employment skill and example")

employment_1_2 = dict.fromkeys(['word1', 'word2'],
                            "sentence of employment skill and example")

employment_list_1 = [x.items() for x in (employment_1_1, employment_1_2)]

cv.scan(employment_list_1)

cv.para("word dates               company	             job title")

employment_2_1 = dict.fromkeys(['word1', 'word2'],
                            "sentence of employment skill and example")

employment_2_2 = dict.fromkeys(['word1', 'word2'],
                            "sentence of employment skill and example")

employment_list_2 = [x.items() for x in (employment_2_1, employment_2_2)]

cv.scan(employment_list_2)


# Education and Certification section
cv.title("Education and Certification")
nothing = document.add_paragraph("", style='List Bullet')
add_hyperlink(nothing, "certificate name", "certificate link")

nothing = document.add_paragraph("", style='List Bullet')

add_hyperlink(nothing, "certificate name", "certificate link")

cv.para("")

cv.para("education years		grades 							school name")

cv.para("education years		grades 							school name")


# This will save file in same file directory as python file
document.save(join(dirname(abspath(__file__)), "NAME.docx"))
