import docx
document = docx.Document()
run = document.add_paragraph().add_run()
'''Apply style'''
style = document.styles['Normal']
font = style.font
font.name = 'MS Gothic'
font.size = docx.shared.Pt(15)
paragraph = document.add_paragraph('Some text\n')
'''Add another sentence to the paragraph'''
sentence = paragraph.add_run('A new line that should have a different font')
'''Then format the sentence'''
sentence.font.name = 'Arial'
sentence.font.size = docx.shared.Pt(10)
paragraph = document.add_paragraph('Some text\n')
document.save("tester.docx")