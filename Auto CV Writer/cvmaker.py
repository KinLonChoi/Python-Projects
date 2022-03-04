import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from hyperlink import add_hyperlink
from os.path import abspath, dirname, join

lines = []
print("Paste job description here and type exit when finished:")
while True:
    line = input()
    if line != 'exit':
        lines.append(line)
    else:
        break
job_desc = ('\n'.join(lines)).lower()
excluded = docx.Document()
excluded.save("Excluded.docx")
document = docx.Document()
document.save("Michael Choi.docx")

# Style for Heading 1
style = document.styles['Heading 1'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Arial'
style.size = Pt(14)
# Font change for heading and titles bugged in docx module asciiTheme also needed to be changed to  Arial.
rFonts = style.element.rPr.rFonts
rFonts.set(qn('w:asciiTheme'), 'Arial')
paragraph_format = document.styles['Heading 1'].paragraph_format
paragraph_format.space_before = Pt(0)

# Style for Heading 2
style = document.styles['Heading 2'].font
style.color.rgb = docx.shared.RGBColor(0, 0, 0)
style.name = 'Arial'
style.size = Pt(13)
rFonts = style.element.rPr.rFonts
rFonts.set(qn('w:asciiTheme'), 'Arial')

# Style for Normal paragraph font
style = document.styles['Normal'].font
style.name = 'Arial'
style.size = Pt(12)

# Style used for information section under heading
styles = document.styles
style = styles.add_style('info', WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles['Normal']
style = document.styles['info'].font
style.name = 'Arial'
style.size = Pt(10)
paragraph_format = document.styles['info'].paragraph_format
paragraph_format.space_after = Pt(0)

# Style for excluded paragraph font
style = excluded.styles['Normal'].font
style.name = 'Arial'
style.size = Pt(12)


class Format:
    def __init__(self, name, github):
        self.name = name
        self.github = github
        # Name and link to GitHub front page is centered
        document.add_heading(name, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph("Wickford, Essex",
                               style='info').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        email = document.add_paragraph("Email: ", style='info')
        add_hyperlink(email, "kinlonchoi@gmail.com", "kinlonchoi@gmail.com")
        document.add_paragraph("Mobile: 07712636191", style='info').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph = document.add_paragraph("GitHub: ", style='info')
        add_hyperlink(paragraph, github, github)
        email.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # The following three functions are made to improve readability of code
    # This is the font used for titles
    def title(self, header):
        document.add_heading(header, level=2)

    # This is the font used for paragraphs
    def para(self, sentence):
        document.add_paragraph(sentence)

    # This is the font used for bullet points
    def bullet(self, sentence):
        document.add_paragraph(sentence, style='List Bullet')

    # Scans job description for keywords
    def scan(self, words_lists):
        for items in words_lists:
            for k, v in items:
                if k in ('front', 'html', 'css', 'web', 'full'):
                    p = document.add_paragraph(v, style='List Bullet')
                    add_hyperlink(p,
                                  "GitHub Website.",
                                  "https://kinlonchoi.github.io/Portfolio/index.html")
                    break
                elif k in ('python', 'js', 'javascript', 'back', 'full'):
                    p = document.add_paragraph(v, style='List Bullet')
                    add_hyperlink(p,
                                  "Python Projects.",
                                  "https://github.com/KinLonChoi/Python-Projects")
                    break
                elif k in job_desc:
                    self.bullet(v)
                    break
                elif k in list(items)[-1]:
                    excluded.add_paragraph(v, style='List Bullet')
                else:
                    continue


cv = Format("Michael Choi", "https://github.com/KinLonChoi")
cv.para("")
cv.para("I am a highly motivated individual with six years of experience in the science industry looking for a career "
        "change to pursue my passion in programming. In my spare time, I have taught myself knowledge of coding in both"
        " front and backend languages. I am keen to learn and develop any skills required to start a career in software"
        " and put existing and newly acquired skills to the test.")

# Skills section add dictionary definition for each key words(search terms) with value(skills) to add as bullet point
cv.title("Skills")


# Words to be searched for in job description
web = dict.fromkeys(['front', 'html', 'css', 'web', 'full'],
                    # if term is found print above as bullet point
                    "Well-versed in web technologies HTML and CSS with responsive designs using bootstrap framework. ")

back = dict.fromkeys(['python', 'js', 'javascript', 'back', 'full'],
                     "Working knowledge of backend languages JavaScript and Python. ")
# Some search terms are shortened to match variations of words e.g. analy will match analytical, analysis etc.
sql = dict.fromkeys(['sql', 'data', 'dbms', 'analy'],
                    "Knowledge of SQL using Database Management System SQLite.")

algo = dict.fromkeys(['algo', 'data structure', 'sdlc', 'lifecycle', 'paradigm'],
                     "Familiar with programming paradigms such as algorithms, data structures and software development "
                     "lifecycles (SDLC).")

time = dict.fromkeys(['pressure', 'result', 'time', 'constraint'],
                     "Accustom to working under time pressure to deliver the results necessary to achieve customer"
                     " requirements in a timely fashion.")

report = dict.fromkeys(['report', 'present', 'analy'],
                       "Experience in data analytics, preparation, and presentation of technical reports according to"
                       " Good Manufacturing Practice (GMP) auditing standards.")

team = dict.fromkeys(['team', 'independent', 'together'],
                     "Effective at working independently during product/method development and as a team to maintain"
                     " high-quality standards throughout the quality department.")

skill_list = [x.items() for x in (web, back, sql, algo, time, report, team)]

cv.scan(skill_list)

# Employment section same as before add dictionary definitions for search terms.
cv.title("Employment")
cv.para("Sept 2016 – Sept 2021               Tate & Lyle PLC	             Laboratory Analyst")

special = dict.fromkeys(['standard', 'test', 'require', 'product'],
                        "Specialised in ensuring products meet the required standards through meticulous testing"
                        " throughout the production process.")

design = dict.fromkeys(['strat' 'design', 'project'],
                       "Designed and implemented appropriate analytic strategies for unique projects.")

collab = dict.fromkeys(['collab', 'team', 'improve', 'meet'],
                       "Collaborated with other departments and developed continuous improvement strategies.")

audit = dict.fromkeys(['report', 'audit', 'feedback'],
                      "Trained and certified in GMP auditing. I audited refinery areas and produced feedback reports"
                      " on non-conformances.")

employment_list_1 = [x.items() for x in (special, design, collab, audit)]

cv.scan(employment_list_1)

cv.para("Feb 2016 – Sept 2016                Tate & Lyle PLC                      Research scientist")

plan = dict.fromkeys(['innovat', 'improve', 'projects', 'plan', 'continuous'],
                     "Planned and conducted experiments for the innovation of new product development and continuous"
                     " improvement projects.")

meet = dict.fromkeys(['meet', 'progress', 'team'],
                     "Exchanged information with colleagues in meetings to maintain steady progress on several"
                     " projects.")

busy = dict.fromkeys(['prior', 'time', 'effici', 'busy'],
                     "Prioritised and allocated time efficiently whilst working on several projects simultaneously.")

employment_list_2 = [x.items() for x in (plan, meet, busy)]

cv.scan(employment_list_2)

# Education and Certification section
cv.title("Education and Certification")
nothing = document.add_paragraph("", style='List Bullet')
add_hyperlink(nothing, "SQL for Data Science by University of California, Davis (Coursera)",
              "https://www.coursera.org/account/accomplishments/certificate/SYPGFMFPK3YA")

nothing = document.add_paragraph("", style='List Bullet')

add_hyperlink(nothing, "Python for Everybody by University of Michigan (Coursera)",
              "https://www.coursera.org/account/accomplishments/specialization/certificate/DK3WDTXK4ND4")

nothing = document.add_paragraph("", style='List Bullet')

add_hyperlink(nothing, "HTML, CSS, and JS for Web Developers by Johns Hopkins University (Coursera)",
              "https://www.coursera.org/account/accomplishments/certificate/ZCNSEFJ3GN6V")

cv.para("")

cv.para("2011-2015		MChem (Hons) in Chemistry (2:1) 							University of Leicester")

cv.para("2003-2011		The Bromfords School, Essex 								"
        "A-Levels in Chemistry (B), Biology (B), and Physics (B) 				"
        "10 GCSEs at A*-C including Chemistry, Maths, and English")

# This will save file in same file directory as python file
document.save(join(dirname(abspath(__file__)), "Michael Choi.docx"))
excluded.save(join(dirname(abspath(__file__)), "Excluded.docx"))
